# -*- coding: utf-8 -*-
"""file_reader 类型路由器单元测试（纯逻辑：内存构造样本，不起服务、不联网）。

覆盖（每类路由 + 降级路径）：
    1.  文本：txt 直读
    2.  文本：超 200KB 截断 + note
    3.  表格：csv → pandas 分析摘要（行列数/列名/数值统计）
    4.  表格：xlsx（openpyxl 现场生成，两个 sheet）
    5.  文档：docx 段落 + 表格（python-docx 现场生成）
    6.  演示文稿：pptx（手工构造 zip + XML）
    7.  文档：pdf（reportlab 现场生成）
    8.  图片：注入假 VLM → 描述进 text；禁用 VLM → 基础信息 + note
    9.  音频：注入假转写 → 文本 + 时长；无转写通道 → note
    10. 音频：whisper 真实冒烟（合成短 wav，全程仅此 1 次真实转写）
    11. 压缩包：清单 + 就地抽取包内文本文件
    12. 二进制：PE 头魔数识别 + strings 片段 + 诚实 note
    13. 降级：损坏的 xlsx → 只降级 note，绝不抛异常

运行：python test_file_reader.py
"""
import io
import os
import struct
import sys
import tempfile
import zipfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import file_reader  # noqa: E402

_checks = 0
_whisper_smoke_done = False  # 全程最多 1 次真实 whisper 转写


def check(name, cond, detail=""):
    global _checks
    _checks += 1
    assert cond, f"{name}: {detail}"
    print(f"✅ {name}")


def make_wav(path, seconds=1.0, rate=16000):
    """合成一段 440Hz 正弦波短 wav（冒烟用，无需真实语音内容）。"""
    import math
    import wave
    n = int(seconds * rate)
    with wave.open(path, "wb") as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(rate)
        frames = b"".join(
            struct.pack("<h", int(12000 * math.sin(2 * math.pi * 440 * i / rate)))
            for i in range(n))
        w.writeframes(frames)


with tempfile.TemporaryDirectory() as tmp:
    # 1. 文本直读
    p = os.path.join(tmp, "note.md")
    with open(p, "w", encoding="utf-8") as f:
        f.write("# 标题\n先生，这是一段测试文本。")
    r = file_reader.read_upload(p, "note.md", vlm_describe=False)
    check("1/13 txt 直读", r["kind"] == "文本" and "测试文本" in r["text"]
          and r["note"] == "" and isinstance(r["meta"], dict), repr(r)[:200])

    # 2. 文本超 200KB 截断 + note
    p = os.path.join(tmp, "big.log")
    with open(p, "w", encoding="utf-8") as f:
        f.write("甲" * (300 * 1024))  # 300KB 字符
    r = file_reader.read_upload(p, "big.log", vlm_describe=False)
    check("2/13 超 200KB 截断", r["meta"].get("truncated") is True
          and "200KB" in r["note"] and len(r["text"]) <= 200 * 1024,
          f"note={r['note']!r} len={len(r['text'])}")

    # 3. csv → pandas 分析摘要
    p = os.path.join(tmp, "sales.csv")
    with open(p, "w", encoding="utf-8") as f:
        f.write("月份,销量\n1月,120\n2月,150\n3月,90\n")
    r = file_reader.read_upload(p, "sales.csv", vlm_describe=False)
    check("3/13 csv 分析摘要", r["kind"] == "表格" and "3 行" in r["text"]
          and "销量" in r["text"] and "数值列统计" in r["text"]
          and r["meta"]["rows"] == 3, repr(r["text"])[:200])

    # 4. xlsx 两个 sheet
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "收入"
    ws.append(["项目", "金额"])
    ws.append(["工资", 10000])
    ws2 = wb.create_sheet("支出")
    ws2.append(["项目", "金额"])
    ws2.append(["房租", 3000])
    p = os.path.join(tmp, "账本.xlsx")
    wb.save(p)
    r = file_reader.read_upload(p, "账本.xlsx", vlm_describe=False)
    check("4/13 xlsx 多 sheet", r["kind"] == "表格" and r["meta"]["sheets"] == ["收入", "支出"]
          and "工作表「收入」" in r["text"] and "房租" in r["text"],
          repr(r["text"])[:200])

    # 5. docx 段落 + 表格
    import docx
    doc = docx.Document()
    doc.add_paragraph("租赁纪要第一段。")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "甲方"
    tbl.cell(0, 1).text = "Nolan"
    tbl.cell(1, 0).text = "租金"
    tbl.cell(1, 1).text = "5000"
    p = os.path.join(tmp, "纪要.docx")
    doc.save(p)
    r = file_reader.read_upload(p, "纪要.docx", vlm_describe=False)
    check("5/13 docx 段落+表格", r["kind"] == "文档" and "租赁纪要" in r["text"]
          and "甲方 | Nolan" in r["text"] and r["meta"]["tables"] == 1,
          repr(r["text"])[:200])

    # 6. pptx 手工构造（zip + slide XML）
    p = os.path.join(tmp, "汇报.pptx")
    slide_tpl = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<p:sld xmlns:p="x" xmlns:a="y"><p:cSld><p:spTree>'
        '<p:sp><p:txBody><a:p><a:r><a:t>%s</a:t></a:r></a:p></p:txBody></p:sp>'
        '</p:spTree></p:cSld></p:sld>')
    with zipfile.ZipFile(p, "w") as zf:
        zf.writestr("ppt/slides/slide1.xml", slide_tpl % "年度汇报标题")
        zf.writestr("ppt/slides/slide2.xml", slide_tpl % "第二页要点 A&lt;B")
    r = file_reader.read_upload(p, "汇报.pptx", vlm_describe=False)
    check("6/13 pptx 按页抽文本", r["kind"] == "演示文稿" and r["meta"]["slides"] == 2
          and "【第 1 页】" in r["text"] and "年度汇报标题" in r["text"]
          and "要点 A<B" in r["text"], repr(r["text"])[:200])

    # 7. pdf（reportlab 现场生成）
    from reportlab.pdfgen import canvas as _rl_canvas
    p = os.path.join(tmp, "报告.pdf")
    c = _rl_canvas.Canvas(p)
    c.drawString(72, 720, "Nolan file reader PDF sample")
    c.save()
    r = file_reader.read_upload(p, "报告.pdf", vlm_describe=False)
    check("7/13 pdf 文字层", r["kind"] == "文档" and "file reader" in r["text"]
          and r["meta"]["pages"] == 1 and r["note"] == "", repr(r["text"])[:200])

    # 8a. 图片：注入假 VLM（开发期真实 VLM 预算走单独脚本，此处零调用）
    from PIL import Image
    p = os.path.join(tmp, "截图.png")
    Image.new("RGB", (320, 200), (30, 30, 40)).save(p)
    r = file_reader.read_upload(
        p, "截图.png", vlm_describe=lambda path: "一张深色背景的测试截图，图中无文字。")
    check("8a/13 图片 VLM 描述", r["kind"] == "图片" and "测试截图" in r["text"]
          and r["meta"]["width"] == 320 and r["meta"].get("described") is True,
          repr(r)[:200])
    # 8b. 图片：禁用 VLM → 基础信息 + 诚实 note
    r = file_reader.read_upload(p, "截图.png", vlm_describe=False)
    check("8b/13 图片无 VLM 降级", r["text"] == "" and "视觉模型" in r["note"]
          and r["meta"]["format"] == "PNG", repr(r["note"])[:200])

    # 9a. 音频：注入假转写
    p = os.path.join(tmp, "语音.wav")
    make_wav(p)
    r = file_reader.read_upload(
        p, "语音.wav", transcribe_fn=lambda path: ("先生你好，这是转写。", 1.0))
    check("9a/13 音频转写", r["kind"] == "音频" and "先生你好" in r["text"]
          and r["meta"]["duration_sec"] == 1.0 and r["note"] == "", repr(r)[:200])
    # 9b. 音频：无转写通道 → note
    r = file_reader.read_upload(p, "语音.wav", transcribe_fn=None)
    check("9b/13 音频无通道降级", r["text"] == "" and "转写通道" in r["note"],
          repr(r["note"])[:200])
    # 9c. 音频：转写抛异常（PyAV 解码失败场景）→ 诚实报错不炸
    def _bad_transcribe(path):
        raise RuntimeError("av decoding failed")
    r = file_reader.read_upload(p, "语音.wav", transcribe_fn=_bad_transcribe)
    check("9c/13 音频解码失败降级", r["text"] == "" and "出错了" in r["note"],
          repr(r["note"])[:200])

    # 10. whisper 真实冒烟（全程仅此 1 次；faster-whisper 不可用时跳过）
    try:
        import faster_whisper  # noqa: F401
        _model = None

        def _real_transcribe(path):
            global _whisper_smoke_done
            nonlocal_model = getattr(_real_transcribe, "_m", None)
            if nonlocal_model is None:
                from faster_whisper import WhisperModel
                nonlocal_model = WhisperModel("small", device="cpu", compute_type="int8")
                _real_transcribe._m = nonlocal_model
            segments, info = nonlocal_model.transcribe(
                path, language="zh", beam_size=1, vad_filter=True)
            return "".join(s.text for s in segments).strip(), float(info.duration)

        r = file_reader.read_upload(p, "语音.wav", transcribe_fn=_real_transcribe)
        _whisper_smoke_done = True
        check("10/13 whisper 真实冒烟（唯一 1 次）",
              r["kind"] == "音频" and r["meta"]["duration_sec"] >= 0.9
              and isinstance(r["text"], str), repr(r)[:200])
    except ImportError:
        print("   （faster-whisper 不可用，10/13 冒烟跳过）")

    # 11. zip：清单 + 就地抽取包内文本
    p = os.path.join(tmp, "打包.zip")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("README.txt", "包内说明：这是清单测试。")
        zf.writestr("data/config.json", '{"key": 42}')
        zf.writestr("binary.dat", b"\x00\x01\x02" * 100)
    with open(p, "wb") as f:
        f.write(buf.getvalue())
    r = file_reader.read_upload(p, "打包.zip", vlm_describe=False)
    check("11/13 zip 清单+抽取", r["kind"] == "压缩包" and r["meta"]["entries"] == 3
          and "README.txt" in r["text"] and "包内说明" in r["text"]
          and '"key": 42' in r["text"] and r["meta"]["extracted_files"] == 2,
          repr(r["text"])[:250])

    # 12. 二进制：PE 头 + strings + 诚实 note
    p = os.path.join(tmp, "tool.exe")
    with open(p, "wb") as f:
        f.write(b"MZ" + b"\x00" * 58 + b"PE\x00\x00" + b"\x00" * 200
                + b"This program cannot be run in DOS mode"
                + b"NolanBinaryTestString12345" + b"\x00" * 100)
    r = file_reader.read_upload(p, "tool.exe", vlm_describe=False)
    check("12/13 二进制魔数+strings", r["kind"] == "二进制"
          and "PE 可执行" in r["meta"].get("detected", "")
          and "NolanBinaryTestString12345" in r["text"]
          and "二进制文件" in r["note"], repr(r)[:250])

    # 13. 降级：损坏的 xlsx（后缀伪装）→ 只降级 note，绝不抛异常
    p = os.path.join(tmp, "坏的.xlsx")
    with open(p, "wb") as f:
        f.write(b"\xde\xad\xbe\xef" * 256)
    r = file_reader.read_upload(p, "坏的.xlsx", vlm_describe=False)
    check("13/13 损坏文件降级不抛异常", r["kind"] == "表格" and r["text"] == ""
          and "出错了" in r["note"], repr(r["note"])[:200])

print(f"\n🎉 file_reader 类型路由器测试全过（{_checks} 项断言；"
      f"whisper 真实冒烟 {'1 次' if _whisper_smoke_done else '跳过'}）")
