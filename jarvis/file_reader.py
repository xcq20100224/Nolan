# -*- coding: utf-8 -*-
"""
file_reader.py —— Nolan 网页版文件拖拽的「全类型阅读引擎」（类型路由器）

第一性原理动机：每一种文件格式都是一条专属的物理通道——
文本是字节序列，表格是行列结构，PDF 是绘制指令流，图片是像素矩阵，
音频是波形，压缩包是容器，二进制是机器码。把 xlsx 当文本读、把图片当
字节丢弃，都是对通道的错配：读出来的不是「内容」，是噪声。
因此本模块按后缀路由到该格式的原生解码器，产出该通道能给出的最诚实结果；
对没有解码通道的格式（未知二进制），不假装读懂——
给出 strings 可读片段 + 文件头魔数识别 + 一句诚实的说明。

统一返回契约（一字不差）：
    {"kind": 类别, "text": 抽取的文本, "meta": {辅助信息}, "note": 给用户的诚实说明}
    kind ∈ 文本/表格/文档/演示文稿/图片/音频/视频/压缩包/二进制
    note 为空串表示通道完全正常；任何单通道失败只降级为 note，绝不抛异常。

额度纪律：
    每类抽取产物上限 200KB（_PRODUCT_CAP）；文本类只读前 200KB；
    PDF 页数 >100 截断并注明；音视频转写文本 >200KB 截断；
    zip 内就地抽取文本类文件 ≤5 个、每个 ≤50KB；二进制 strings ≤100 条。
"""

import base64
import io
import json
import os
import re
import zipfile

# ---------------------------------------------------------------------------
# 常量与额度
# ---------------------------------------------------------------------------

_PRODUCT_CAP = 200 * 1024          # 每类抽取产物上限 200KB（字符数近似）
_TEXT_READ_LIMIT = 200 * 1024      # 文本类只读前 200KB
_PDF_MAX_PAGES = 100               # PDF 页数上限（超出截断并注明）
_ZIP_TEXT_MAX_FILES = 5            # zip 内就地抽取的文本文件数上限
_ZIP_TEXT_MAX_BYTES = 50 * 1024    # zip 内单个文本文件抽取上限
_STRINGS_MAX = 100                 # 二进制可读串扫描条数上限
_STRINGS_MIN_LEN = 6               # 可读串最小长度

_CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "llm_config.json")

# 纯文本直读通道：字节序列即内容，UTF-8 容错解码即可
_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".py", ".js", ".ts", ".jsx", ".tsx",
    ".json", ".jsonl", ".log", ".ini", ".cfg", ".conf", ".yaml", ".yml",
    ".xml", ".toml", ".sql", ".sh", ".bat", ".ps1", ".css", ".scss",
    ".less", ".env", ".gitignore", ".properties", ".rst", ".tex",
}
# 结构化文本：字节序列之下还有一层结构，直读会丢信息——走专属通道
_CSV_EXTS = {".csv", ".tsv"}       # 行列结构 → pandas 附加分析摘要
_HTML_EXTS = {".html", ".htm"}     # 标签噪声 → bs4 去标签
_XLSX_EXTS = {".xlsx", ".xlsm"}    # 多 sheet 行列结构 → pandas/openpyxl
_DOCX_EXTS = {".docx"}             # 段落 + 表格 → python-docx
_PPTX_EXTS = {".pptx"}             # 幻灯片文本框 → zipfile + XML 手工解析
_PDF_EXTS = {".pdf"}               # 绘制指令流 → pypdf 文字层
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp"}
_AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".ogg", ".flac"}
_VIDEO_EXTS = {".mp4", ".mov", ".mkv"}
_ZIP_EXTS = {".zip"}

# 常见文件头魔数表：未知二进制的第一条线索——头几个字节就是格式指纹
_MAGIC_TABLE = [
    (b"MZ", "PE 可执行程序（Windows exe/dll）"),
    (b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "OLE 复合文档（老版 Office doc/xls/ppt）"),
    (b"SQLite format 3\x00", "SQLite 数据库文件"),
    (b"\x7fELF", "ELF 可执行程序（Linux）"),
    (b"PK\x03\x04", "ZIP 容器（或改后缀的 Office/jar/apk）"),
    (b"Rar!\x1a\x07", "RAR 压缩包"),
    (b"7z\xbc\xaf\x27\x1c", "7z 压缩包"),
    (b"\x1f\x8b", "gzip 压缩数据"),
    (b"%PDF", "PDF 文档（后缀可能被改）"),
    (b"\x89PNG\r\n\x1a\n", "PNG 图片（后缀可能被改）"),
    (b"\xff\xd8\xff", "JPEG 图片（后缀可能被改）"),
    (b"GIF8", "GIF 图片（后缀可能被改）"),
    (b"ID3", "MP3 音频（含 ID3 标签）"),
    (b"OggS", "OGG 音视频容器"),
    (b"RIFF", "RIFF 媒体（wav/avi 等）"),
    (b"\x00\x00\x00", "ISO BMFF 媒体（mp4/mov 等，若第 5-8 字节为 ftyp）"),
]


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------

def _cap(text: str, meta: dict, limit: int = _PRODUCT_CAP) -> str:
    """产物封顶：超出 limit 字符截断并在 meta 里标注。"""
    text = text or ""
    if len(text) > limit:
        meta["truncated"] = True
        return text[:limit]
    return text


def _identify_magic(head: bytes) -> str:
    """按文件头魔数识别真实格式；识别不了返回空串。"""
    for magic, label in _MAGIC_TABLE:
        if head.startswith(magic):
            if magic == b"\x00\x00\x00" and head[4:8] != b"ftyp":
                continue  # 前 3 字节为零但非 ftyp，不判 ISO BMFF
            return label
    return ""


def _scan_strings(data: bytes, limit: int = _STRINGS_MAX) -> list:
    """strings 式可读片段扫描：长度 ≥6 的 ASCII/UTF-8 可读串，最多 limit 条。"""
    out = []
    for m in re.finditer(rb"[\x20-\x7e]{%d,}" % _STRINGS_MIN_LEN, data):
        s = m.group().decode("ascii", errors="replace")
        out.append(s)
        if len(out) >= limit:
            break
    return out


def _df_summary(df, name: str, head_rows: int = 5) -> str:
    """一个 DataFrame 的分析摘要：行列数、列名、前 5 行、数值列 describe。

    第一性原理：表格的内容不是字，是结构——只抽字会把「1000 行销售数据」
    读成一堆数字碎片；给出形状 + 列名 + 样本 + 统计，大脑才能回答
    「这个表有多少行、平均是多少」这类问题。
    """
    parts = [
        f"【{name}】{df.shape[0]} 行 × {df.shape[1]} 列",
        "列名：" + "、".join(str(c) for c in df.columns),
        f"前 {min(head_rows, len(df))} 行：",
        df.head(head_rows).to_string(index=False),
    ]
    num = df.select_dtypes(include="number")
    if num.shape[1] > 0 and len(df) > 0:
        parts.append("数值列统计：")
        parts.append(num.describe().to_string())
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 各类型专属通道（每个函数只负责一种格式；失败抛异常由路由器统一降级）
# ---------------------------------------------------------------------------

def _read_text(path: str, meta: dict) -> tuple:
    """文本/代码/配置：直读前 200KB，UTF-8 容错解码。"""
    with open(path, "rb") as f:
        raw = f.read(_TEXT_READ_LIMIT + 1)
    if len(raw) > _TEXT_READ_LIMIT:
        raw = raw[:_TEXT_READ_LIMIT]
        meta["truncated"] = True
        meta["note_extra"] = f"文本超过 200KB，只读取了前 200KB。"
    return raw.decode("utf-8", errors="replace"), ""


def _read_csv(path: str, meta: dict) -> tuple:
    """csv/tsv：pandas 读入 + 结构分析摘要（行列数、列名、前 5 行、数值列统计）。"""
    import pandas as pd
    sep = "\t" if path.lower().endswith(".tsv") else ","
    df = pd.read_csv(path, sep=sep, encoding="utf-8",
                     encoding_errors="replace", on_bad_lines="skip")
    meta["rows"], meta["cols"] = int(df.shape[0]), int(df.shape[1])
    meta["columns"] = [str(c) for c in df.columns]
    return _df_summary(df, "表格"), ""


def _read_html(path: str, meta: dict) -> tuple:
    """html：bs4 去标签留正文（脚本/样式剔除），直读前 200KB。"""
    from bs4 import BeautifulSoup
    with open(path, "rb") as f:
        raw = f.read(_TEXT_READ_LIMIT + 1)
    if len(raw) > _TEXT_READ_LIMIT:
        raw = raw[:_TEXT_READ_LIMIT]
        meta["truncated"] = True
    soup = BeautifulSoup(raw.decode("utf-8", errors="replace"), "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    text = soup.get_text("\n")
    # 压缩连续空行：网页正文抽出后常带大量空行
    text = re.sub(r"\n\s*\n+", "\n\n", text).strip()
    return text, ""


def _read_xlsx(path: str, meta: dict) -> tuple:
    """xlsx/xlsm：sheet 名单 + 每个 sheet 的行列数、列名、前 5 行、数值列统计。"""
    import pandas as pd
    sheets = pd.read_excel(path, sheet_name=None, engine="openpyxl")
    meta["sheets"] = list(sheets.keys())
    parts = [f"共 {len(sheets)} 个工作表：" + "、".join(sheets.keys())]
    for name, df in sheets.items():
        parts.append("")
        parts.append(_df_summary(df, f"工作表「{name}」"))
    return "\n".join(parts), ""


def _read_docx(path: str, meta: dict) -> tuple:
    """docx：段落 + 表格（旧版只抽段落，表格整块丢失——合同/报价单的核心
    常在表格里，必须补上）。"""
    import docx
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    meta["paragraphs"] = len(parts)
    meta["tables"] = len(document.tables)
    for i, table in enumerate(document.tables, 1):
        parts.append(f"\n【表格 {i}】")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts).strip(), ""


def _read_pptx(path: str, meta: dict) -> tuple:
    """pptx：无 python-pptx，zipfile + XML 手工解析——
    pptx 本质是 zip 容器，每页幻灯片是 ppt/slides/slideN.xml，
    文本框文字在 <a:t> 标签里，按页抽出即可。"""
    texts = {}
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            m = re.match(r"ppt/slides/slide(\d+)\.xml$", name)
            if not m:
                continue
            xml = zf.read(name).decode("utf-8", errors="replace")
            runs = re.findall(r"<a:t>(.*?)</a:t>", xml, re.DOTALL)
            # XML 实体反转义（&amp; &lt; &gt; 等）
            runs = [r.replace("&lt;", "<").replace("&gt;", ">")
                    .replace("&amp;", "&").replace("&quot;", '"')
                    .replace("&apos;", "'") for r in runs]
            texts[int(m.group(1))] = "".join(runs).strip()
    meta["slides"] = len(texts)
    if not texts:
        return "", "PPTX 里没有抽到任何文本（可能全部为图片型幻灯片）。"
    parts = []
    for n in sorted(texts):
        if texts[n]:
            parts.append(f"【第 {n} 页】\n{texts[n]}")
    return "\n\n".join(parts), ""


def _read_pdf(path: str, meta: dict) -> tuple:
    """pdf：pypdf 抽文字层，全页（>100 页截断注明）。
    扫描件没有文字层——抽不出文本时诚实声明，绝不假装读了。"""
    from pypdf import PdfReader
    reader = PdfReader(path)
    total = len(reader.pages)
    meta["pages"] = total
    if reader.is_encrypted:
        return "", "PDF 已加密，无法读取内容。"
    pages = reader.pages[:_PDF_MAX_PAGES]
    parts = []
    for page in pages:
        parts.append(page.extract_text() or "")
    text = "\n".join(p for p in parts if p).strip()
    if total > _PDF_MAX_PAGES:
        meta["truncated"] = True
        return text, f"PDF 共 {total} 页，只抽取了前 {_PDF_MAX_PAGES} 页。"
    if total > 0 and len(text) < 20:
        return text, "扫描版 PDF，无文本层，无法抽取文字内容。"
    return text, ""


def _read_image(path: str, meta: dict, vlm_describe) -> tuple:
    """图片：PIL 读基础信息（尺寸/格式），视觉模型做内容解读
    （画面内容、图中文字、关键元素，中文 ≤200 字）。
    VLM 不可用时只返回基础信息 + 诚实 note。"""
    from PIL import Image
    with Image.open(path) as img:
        meta["width"], meta["height"] = img.size
        meta["format"] = img.format or ""
    if vlm_describe is None:
        return "", ("这是一张图片（%s，%d×%d）。视觉模型暂不可用，"
                    "只能提供尺寸与格式信息。" % (
                        meta["format"] or "未知格式",
                        meta["width"], meta["height"]))
    try:
        desc = (vlm_describe(path) or "").strip()
    except Exception as e:
        return "", ("这是一张图片（%s，%d×%d）。视觉模型解读失败了（%s），"
                    "只能提供尺寸与格式信息。" % (
                        meta["format"] or "未知格式",
                        meta["width"], meta["height"], e))
    if not desc:
        return "", "视觉模型没有返回有效描述，只能提供尺寸与格式信息。"
    meta["described"] = True
    return desc, ""


def _read_av(path: str, meta: dict, transcribe_fn, kind: str) -> tuple:
    """音频/视频：复用常驻 whisper 转写（由 server 注入 transcribe_fn）。
    PyAV 解码失败的格式诚实报错；>10 分钟的转写文本截断到 200KB。"""
    if transcribe_fn is None:
        return "", f"这是一个{kind}文件，但语音转写通道当前不可用，无法读取内容。"
    text, duration = transcribe_fn(path)  # 解码失败抛异常 → 路由器降级
    meta["duration_sec"] = round(float(duration or 0.0), 1)
    meta["transcribed"] = bool(text)
    if not text:
        return "", (f"{kind}时长约 {meta['duration_sec']} 秒，"
                    "未识别到有效语音内容（可能是纯音乐或静音）。")
    if meta["duration_sec"] > 600 and len(text) > _PRODUCT_CAP:
        meta["truncated"] = True
        return text[:_PRODUCT_CAP], (f"{kind}超过 10 分钟，"
                                     "转写文本已截断到前 200KB。")
    return text, ""


def _read_zip(path: str, meta: dict) -> tuple:
    """zip 压缩包：列出内容清单（路径+大小），并就地抽取其中文本类文件
    （≤5 个、每个 ≤50KB）——压缩包是容器，「内容」= 清单 + 成员的内容。"""
    parts = []
    extracted = 0
    with zipfile.ZipFile(path) as zf:
        infos = zf.infolist()
        meta["entries"] = len(infos)
        parts.append(f"压缩包内共 {len(infos)} 个条目：")
        for info in infos[:50]:  # 清单最多列 50 条，防巨型包刷屏
            if info.is_dir():
                continue
            size_kb = info.file_size / 1024
            parts.append(f"  {info.filename}（{size_kb:.1f}KB）")
        if len(infos) > 50:
            parts.append(f"  ……（其余 {len(infos) - 50} 条略）")
        for info in infos:
            if extracted >= _ZIP_TEXT_MAX_FILES:
                break
            if info.is_dir():
                continue
            ext = os.path.splitext(info.filename)[1].lower()
            if ext not in (_TEXT_EXTS | _CSV_EXTS):
                continue
            if info.file_size > _ZIP_TEXT_MAX_BYTES:
                continue
            try:
                raw = zf.read(info.filename)
            except Exception:
                continue
            extracted += 1
            parts.append(f"\n【包内文件：{info.filename}】")
            parts.append(raw.decode("utf-8", errors="replace"))
    meta["extracted_files"] = extracted
    return "\n".join(parts), ""


def _read_binary(path: str, meta: dict) -> tuple:
    """其他一切二进制：文件头魔数识别 + strings 式可读片段扫描。
    没有专属解码通道的格式，绝不假装读懂——给出可识别的线索，
    附一句诚实说明。"""
    with open(path, "rb") as f:
        head = f.read(64 * 1024)  # 读前 64KB 扫描足够（strings 与魔数都在头部）
    size = os.path.getsize(path)
    meta["size_bytes"] = size
    label = _identify_magic(head)
    if label:
        meta["detected"] = label
    strings = _scan_strings(head)
    meta["strings_found"] = len(strings)
    parts = [f"文件大小：{size / 1024:.1f}KB"]
    if label:
        parts.append(f"格式识别：{label}")
    if strings:
        parts.append("可读片段：")
        parts.extend(f"  {s}" for s in strings)
    note = "这是一个二进制文件，无法直接阅读；以下是可识别的信息。"
    return "\n".join(parts), note


# ---------------------------------------------------------------------------
# VLM 默认通道（图片解读；server 可注入自己的实现）
# ---------------------------------------------------------------------------

_VLM_FALLBACK_MODEL = "glm-4v-flash"
_VLM_IMAGE_MAX_SIDE = 1280         # 发给 VLM 的图片最长边（省 token、防超限）

_IMAGE_PROMPT = (
    "请用中文客观描述这张图片，不超过 200 字。包括："
    "1) 画面主要内容是什么；2) 图中出现的文字（如有，尽量完整抄录）；"
    "3) 关键元素与数据（如截图界面、图表读数、文档标题）。"
    "只输出描述本身，不要解释、不要前缀。"
)


def _vlm_call(image_b64: str, mime: str, model: str, extra: dict) -> str:
    """单次视觉模型调用（OpenAI 兼容 image_url data URI，与 eyes.py 同写法）。"""
    import httpx
    with open(_CONFIG_PATH, "r", encoding="utf-8") as f:
        cfg = json.load(f)
    url = str(cfg.get("vision_base_url") or cfg["base_url"]).rstrip("/") \
        + "/chat/completions"
    headers = {
        "Authorization": "Bearer " + str(cfg.get("vision_api_key") or cfg["api_key"]),
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": _IMAGE_PROMPT},
                {"type": "image_url",
                 "image_url": {"url": f"data:{mime};base64,{image_b64}"}},
            ],
        }],
        "temperature": 0.1,
    }
    payload.update(extra)
    with httpx.Client(timeout=60.0) as client:
        resp = client.post(url, headers=headers, json=payload)
        resp.raise_for_status()
        data = resp.json()
    return data["choices"][0]["message"]["content"]


def default_vlm_describe(path: str) -> str:
    """默认图片解读通道：缩放编码 → glm-4.5v（配置驱动），
    主模型失败自动降级 glm-4v-flash 一次（与 eyes.py 同一策略）。"""
    from PIL import Image
    with Image.open(path) as img:
        if max(img.size) > _VLM_IMAGE_MAX_SIDE:
            ratio = _VLM_IMAGE_MAX_SIDE / max(img.size)
            img = img.resize((max(1, round(img.width * ratio)),
                              max(1, round(img.height * ratio))),
                             Image.LANCZOS)
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG", quality=85)
    b64 = base64.b64encode(buf.getvalue()).decode("ascii")
    with open(_CONFIG_PATH, "r", encoding="utf-8") as f:
        cfg = json.load(f)
    model = str(cfg.get("vision_model") or "glm-4.5v").strip() or "glm-4.5v"
    extra_raw = cfg.get("vision_extra_body", '{"thinking": {"type": "disabled"}}')
    try:
        extra = extra_raw if isinstance(extra_raw, dict) else json.loads(str(extra_raw))
        if not isinstance(extra, dict):
            extra = {}
    except (ValueError, TypeError):
        extra = {}
    try:
        return _vlm_call(b64, "image/jpeg", model, extra).strip()
    except Exception:
        if model == _VLM_FALLBACK_MODEL:
            raise
        # 降级链不带 extra_body：thinking 参数是 glm-4.5 系专属，带上可能直接报错
        return _vlm_call(b64, "image/jpeg", _VLM_FALLBACK_MODEL, {}).strip()


# ---------------------------------------------------------------------------
# 类型路由器（唯一入口）
# ---------------------------------------------------------------------------

def read_upload(path: str, name: str, *,
                transcribe_fn=None, vlm_describe=None) -> dict:
    """
    按后缀把文件路由到专属通道，返回统一 dict：
        {"kind": 类别, "text": 抽取的文本, "meta": {辅助信息}, "note": 诚实说明}

    参数：
        path           落盘后的文件路径
        name           原始文件名（取后缀用）
        transcribe_fn  音视频转写注入：fn(path) -> (text, duration_sec)；
                       解码失败抛异常；None 表示转写通道不可用
        vlm_describe   图片解读注入：fn(path) -> str；None 时用默认 VLM 通道
                       （llm_config.json 不可读时自动退化为基础信息 + note）

    失败纪律：任何单通道异常只降级为 note + 尽量保留的 meta，绝不抛异常——
    上传通道的底线是「收得下、说得清」，不是「读得出」。
    """
    ext = os.path.splitext(name or "")[1].lower()
    meta: dict = {"ext": ext}
    kind, text, note = "二进制", "", ""
    try:
        if ext in _TEXT_EXTS:
            kind = "文本"
            text, note = _read_text(path, meta)
        elif ext in _CSV_EXTS:
            kind = "表格"
            text, note = _read_csv(path, meta)
        elif ext in _HTML_EXTS:
            kind = "文本"
            text, note = _read_html(path, meta)
        elif ext in _XLSX_EXTS:
            kind = "表格"
            text, note = _read_xlsx(path, meta)
        elif ext in _DOCX_EXTS:
            kind = "文档"
            text, note = _read_docx(path, meta)
        elif ext in _PPTX_EXTS:
            kind = "演示文稿"
            text, note = _read_pptx(path, meta)
        elif ext in _PDF_EXTS:
            kind = "文档"
            text, note = _read_pdf(path, meta)
        elif ext in _IMAGE_EXTS:
            kind = "图片"
            # vlm_describe 未提供（None）时尝试默认 VLM 通道，配置不可读则退化；
            # 显式传 False 表示「禁用 VLM」（测试/离线场景），只给基础信息 + note
            fn = vlm_describe
            if fn is False:
                fn = None
            elif fn is None:
                fn = default_vlm_describe if os.path.isfile(_CONFIG_PATH) else None
            text, note = _read_image(path, meta, fn)
        elif ext in _AUDIO_EXTS:
            kind = "音频"
            text, note = _read_av(path, meta, transcribe_fn, kind)
        elif ext in _VIDEO_EXTS:
            kind = "视频"
            text, note = _read_av(path, meta, transcribe_fn, kind)
        elif ext in _ZIP_EXTS:
            kind = "压缩包"
            text, note = _read_zip(path, meta)
        else:
            kind = "二进制"
            text, note = _read_binary(path, meta)
    except Exception as e:
        # 单通道失败：降级为诚实 note，产物为空，绝不把异常透出给 HTTP 层
        note = f"读取这个文件时出错了（{type(e).__name__}：{e}），无法抽取内容。"
        text = ""
    # 通道内部 note_extra（如文本超 200KB）与显式 note 合并
    extra_note = meta.pop("note_extra", "")
    if extra_note:
        note = (note + " " + extra_note).strip() if note else extra_note
    text = _cap(text, meta)
    return {"kind": kind, "text": text, "meta": meta, "note": note}
