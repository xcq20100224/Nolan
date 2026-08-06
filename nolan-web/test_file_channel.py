# -*- coding: utf-8 -*-
"""文件通道单元测试（纯逻辑，不起服务、不写真实 jarvis/files/uploads）。

全类型升级版契约：_save_upload 返回 (存储名, 抽取结果 dict)，
dict = {"kind", "text", "meta", "note"}（jarvis/file_reader.py 类型路由器）；
任何格式都收（无类型白名单），任何单类失败只降级 note。

覆盖：
    1.  文本抽取：txt 正常读取
    2.  文本抽取：超过 200KB 截断 + note
    3.  PDF 抽取：reportlab 内存构造样本 → kind=文档 + extracted.txt 落盘
    4.  DOCX 抽取：段落 + 表格
    5.  文件名净化：路径剥离、危险字符替换、空名兜底
    6.  大小上限：>50MB 拒绝；>200KB 文本文件正常截断抽取
    7.  未知二进制（exe）：不再拒收——二进制通道 + 诚实 note，文件留柜
    8.  上传目录防护：uploads 不在 files 内时拒绝写入
    9.  files_list 结构：字段齐全、kind 合法、mtime 倒序、内部文件已排除
    10. 路径穿越拒绝：'..' / 绝对路径 / 盘符一律 None；合法相对路径放行
    11. xlsx/zip 全类型通道经 _save_upload 端到端走通

运行：python test_file_channel.py
"""
import importlib.util
import io
import os
import sys
import tempfile
import zipfile

# 动态加载 server.py（不启动 HTTP 服务，只取纯函数，与 test_bargein_echo 同一模式）
spec = importlib.util.spec_from_file_location(
    "nolan_server", os.path.join(os.path.dirname(__file__), "server.py"))
server = importlib.util.module_from_spec(spec)
sys.modules["nolan_server"] = server
spec.loader.exec_module(server)

_checks = 0


def check(name, cond, detail=""):
    global _checks
    _checks += 1
    assert cond, f"{name}: {detail}"
    print(f"✅ {name}")


def make_sample_pdf(path):
    """用 reportlab 在内存里构造一页带文本的 PDF 样本。"""
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(path)
    c.drawString(72, 720, "Hello Nolan file channel PDF sample")
    c.save()


def make_sample_docx(path):
    """用 python-docx 在内存里构造带段落 + 表格的 DOCX 样本。"""
    import docx
    doc = docx.Document()
    doc.add_paragraph("第一段：Nolan 文件通道测试。")
    doc.add_paragraph("第二段：拖拽阅读验收。")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "项目"
    tbl.cell(0, 1).text = "全类型升级"
    doc.save(path)


with tempfile.TemporaryDirectory() as tmp:
    uploads = os.path.join(tmp, "uploads")

    # 1. 文本抽取：txt 正常读取（经 file_reader 路由器）
    txt_path = os.path.join(tmp, "note.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("先生，这是一段测试文本。")
    r = server.file_reader.read_upload(txt_path, "note.txt", vlm_describe=False)
    check("1/11 txt 文本抽取",
          r["kind"] == "文本" and r["text"] == "先生，这是一段测试文本。"
          and r["note"] == "" and isinstance(r["meta"], dict),
          repr(r)[:200])

    # 2. 文本抽取：超过 200KB 截断 + note
    big_path = os.path.join(tmp, "big.txt")
    with open(big_path, "w", encoding="utf-8") as f:
        f.write("甲" * (300 * 1024))  # 300KB 汉字
    r = server.file_reader.read_upload(big_path, "big.txt", vlm_describe=False)
    check("2/11 超 200KB 文本截断",
          r["meta"].get("truncated") is True and "200KB" in r["note"]
          and len(r["text"]) <= 200 * 1024,
          f"note={r['note']!r} chars={len(r['text'])}")

    # 3. PDF 抽取（内存构造样本，走完整 _save_upload 链路 + extracted.txt 落盘）
    pdf_path = os.path.join(tmp, "sample.pdf")
    make_sample_pdf(pdf_path)
    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()
    stored, r = server._save_upload(
        "报告.pdf", pdf_bytes, uploads_dir=uploads, files_dir=tmp)
    check("3a/11 PDF 抽取", r["kind"] == "文档" and "Nolan file channel" in r["text"]
          and stored.endswith("_报告.pdf"), f"stored={stored!r} text={r['text'][:60]!r}")
    extracted = os.path.join(uploads, stored + ".extracted.txt")
    check("3b/11 抽取全文落盘",
          os.path.isfile(extracted)
          and "Nolan file channel" in open(extracted, encoding="utf-8").read(),
          extracted)

    # 4. DOCX 抽取：段落 + 表格（旧版漏表格，新版必须补上）
    docx_path = os.path.join(tmp, "sample.docx")
    make_sample_docx(docx_path)
    with open(docx_path, "rb") as f:
        docx_bytes = f.read()
    stored, r = server._save_upload(
        "纪要.docx", docx_bytes, uploads_dir=uploads, files_dir=tmp)
    check("4/11 DOCX 段落+表格",
          r["kind"] == "文档" and "文件通道测试" in r["text"]
          and "拖拽阅读验收" in r["text"] and "项目 | 全类型升级" in r["text"]
          and r["meta"]["tables"] == 1, repr(r["text"])[:200])

    # 5. 文件名净化
    check("5a/11 净化：剥离路径",
          server._sanitize_filename("../../../etc/passwd") == "passwd")
    check("5b/11 净化：Windows 反斜杠路径",
          server._sanitize_filename("..\\..\\evil.txt") == "evil.txt")
    check("5c/11 净化：危险字符替换",
          server._sanitize_filename('a<b>:"|?*.txt') == "a_b______.txt",
          server._sanitize_filename('a<b>:"|?*.txt'))
    check("5d/11 净化：空名与纯点兜底",
          server._sanitize_filename("..") == "file"
          and server._sanitize_filename("") == "file")
    check("5e/11 净化：中文名保留",
          server._sanitize_filename("租房合同 注意事项.pdf") == "租房合同_注意事项.pdf",
          server._sanitize_filename("租房合同 注意事项.pdf"))

    # 6. 大小上限（50MB）
    try:
        server._save_upload("big.txt", b"x" * (server._UPLOAD_MAX_BYTES + 1),
                            uploads_dir=uploads, files_dir=tmp)
        raise AssertionError("6/11 超限文件未被拒绝")
    except ValueError as e:
        check("6/11 超 50MB 拒绝", "上限" in str(e) and "50MB" in str(e), str(e))

    # 7. 未知二进制（exe）：全类型升级后不再拒收——二进制通道 + 诚实 note，文件留柜
    stored, r = server._save_upload("tool.exe", b"MZ" + b"\0" * 100,
                                    uploads_dir=uploads, files_dir=tmp)
    check("7/11 exe 走二进制通道且留柜",
          r["kind"] == "二进制" and "二进制文件" in r["note"]
          and "PE 可执行" in r["meta"].get("detected", "")
          and os.path.isfile(os.path.join(uploads, stored)),
          f"kind={r['kind']!r} note={r['note']!r}")

    # 8. 上传目录防护：uploads 不在 files 内时拒绝写入
    with tempfile.TemporaryDirectory() as outside:
        try:
            server._save_upload("a.txt", b"hello",
                                uploads_dir=os.path.join(outside, "up"),
                                files_dir=tmp)
            raise AssertionError("8/11 越界目录未被拒绝")
        except ValueError as e:
            check("8/11 uploads 越界拒绝写入", "不在文件柜内" in str(e), str(e))

    # 9. files_list 结构（只读扫描真实 jarvis/files/，不写入）
    payload = server._files_list_payload()
    files = payload.get("files")
    check("9a/11 files_list 返回列表", isinstance(files, list))
    if files:
        item = files[0]
        check("9b/11 files_list 字段齐全",
              all(k in item for k in ("name", "size", "mtime", "kind")),
              f"keys={sorted(item.keys())}")
        check("9c/11 files_list kind 合法",
              all(f["kind"] in ("文档", "图片", "表格", "音频", "其他") for f in files))
        mtimes = [f["mtime"] for f in files]
        check("9d/11 files_list mtime 倒序", mtimes == sorted(mtimes, reverse=True))
        check("9e/11 files_list 排除内部文件",
              not any(f["name"].startswith("tts_cache/")
                      or f["name"] == "server.pid"
                      or f["name"].endswith(".log") for f in files),
              [f["name"] for f in files][:5])
    else:
        print("   （文件柜当前为空，结构子项跳过）")

    # 10. 路径穿越拒绝
    check("10a/11 穿越 '..'", server._resolve_files_path("../brain.py") is None)
    check("10b/11 穿越反斜杠", server._resolve_files_path("..\\brain.py") is None)
    check("10c/11 盘符拒绝", server._resolve_files_path("C:\\Windows\\x.dll") is None)
    check("10d/11 空名拒绝", server._resolve_files_path("") is None)
    ok_path = server._resolve_files_path("uploads/20260101-000000_a.txt")
    check("10e/11 合法子目录放行",
          ok_path is not None
          and os.path.commonpath([ok_path, server._FILES_DIR]) == server._FILES_DIR,
          repr(ok_path))

    # 11a. xlsx 端到端（openpyxl 现场生成）
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "数据"
    ws.append(["名称", "数量"])
    ws.append(["苹果", 7])
    buf = io.BytesIO()
    wb.save(buf)
    stored, r = server._save_upload(
        "清单.xlsx", buf.getvalue(), uploads_dir=uploads, files_dir=tmp)
    check("11a/11 xlsx 端到端",
          r["kind"] == "表格" and "苹果" in r["text"]
          and r["meta"]["sheets"] == ["数据"], repr(r["text"])[:150])

    # 11b. zip 端到端（内存打包）
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("说明.txt", "端到端 zip 测试内容。")
    stored, r = server._save_upload(
        "资料.zip", buf.getvalue(), uploads_dir=uploads, files_dir=tmp)
    check("11b/11 zip 端到端",
          r["kind"] == "压缩包" and "说明.txt" in r["text"]
          and "端到端 zip 测试内容" in r["text"], repr(r["text"])[:200])

print(f"\n🎉 文件通道单元测试全过（{_checks} 项断言）")
